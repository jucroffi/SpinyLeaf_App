# %%
import os

from llama_cpp import Llama

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from pathlib import Path

import pandas as pd
import re

from pathlib import Path
import re


MASTER_BIB = {

    ("barrie", "2023"): "Barrie, H., McDougall, K., Miller, K., & Faulkner, D. (2023). The social value of public spaces in mixed-use high-rise buildings. *Buildings and Cities, 4*(1), 669–689. https://doi.org/10.5334/bc.339",
    ("bee", "2016"): "Bee, C., & Im, J. (2016). Neighborhood design and community building. *Journal of Urban Design, 21*(1), 1–29. https://doi.org/10.1080/13574809.2015.1106920",
    ("aw", "2016"): "Aw, S. B., & Lim, P. I. (2016). The provision of vertical social pockets for better social interaction in high-rise living. *Planning Malaysia: Journal of the Malaysian Institute of Planners, Special Issue IV*, 163–180.",
    ("jordan", "2023"): "Jordan, H. (2023). The impact of social connections on health and wellbeing. *Health & Social Care in the Community, 31*(1), 85–95. https://doi.org/10.1111/hsc.13889",
    ("kearns", "2012"): "Kearns, A., Whitley, E., Tannahill, C., & Ellaway, A. (2012). Loneliness, social relations and health and well-being in deprived communities. *Psychology, Health & Medicine, 17*(3), 374–385. https://doi.org/10.1080/13548506.2011.608805",
    ("umberson", "2010"): "Umberson, D., & Montez, J. K. (2010). Social relationships and health: A flashpoint for health policy. *Journal of Health and Social Behavior, 51*(Suppl), S54–S66. https://doi.org/10.1177/0022146510383501",
    ("weijs-perree", "2015"): "Weijs-Perrée, M., van den Berg, P., Arentze, T., & Kemperman, A. (2015). Factors influencing social satisfaction and loneliness: A path analysis. *Journal of Transport & Health, 2*(4), 451–462. https://doi.org/10.1016/j.jth.2015.07.002",
    ("williams", "2005"): "Williams, K. (2005). Spatial planning, urban form and sustainable transport: An evidence-based approach. *Urban Studies, 42*(7), 1145–1166. https://doi.org/10.1080/00420980500143646",
    ("williams-cohousing", "2005"): "Williams, J. (2005). Designing neighbourhoods for social interaction: The case of cohousing. *Journal of Urban Design, 10*(2), 195–227. https://doi.org/10.1080/13574800500086998",
    ("barros", "2019"): "Barros, P., Ng Fat, L., Garcia, L. M. T., Slovic, A. D., Thomopoulos, N., de Sá, T. H., Morais, P., et al. (2019). Social consequences and mental health outcomes of living in high-rise residential buildings: A systematic review. *Cities, 93*, 263–272. https://doi.org/10.1016/j.cities.2019.05.015",
    ("kleeman", "2023"): "Kleeman, A., Giles-Corti, B., Gunn, L., Hooper, P., & Foster, S. (2023). The impact of the design and quality of communal areas in apartment buildings on residents’ neighbouring and loneliness. *Cities, 133*, 104126. https://doi.org/10.1016/j.cities.2022.104126",
    ("leavell", "2019"): "Leavell, M. A., Leiferman, J. A., Gascon, M., Braddick, F., Gonzalez, J. C., & Litt, J. S. (2019). Nature-based social prescribing in urban settings to improve social connectedness and mental well-being: A review. *Current Environmental Health Reports, 6*(4), 297–308. https://doi.org/10.1007/s40572-019-00251-7",
    ("nguyen", "2024"): "Nguyen, L. P., Van Den Berg, P. E., Kemperman, A. D., & Mohammadi, M. (2024). Social impacts of living in high-rise apartment buildings: The effects of buildings and neighborhoods. *Journal of Urban Affairs*, 1–22. https://doi.org/10.1080/07352166.2024.2311165",
    ("nguyen", "2025"): "Nguyen, L., Van Den Berg, P., Kemperman, A., & Mohammadi, M. (2025). How does the layout of indoor communal spaces in low-income high-rise apartment buildings impact the social interactions between residents? *Cities & Health*, 1–24. https://doi.org/10.1080/23748834.2025.2509739",

    
    ("altomonte", "2020"): "Altomonte, S., Allen, J., Bluyssen, P. M., Brager, G., Heschong, L., Loder, A., Schiavon, S., et al. (2020). Ten questions concerning well-being in the built environment. *Building and Environment, 180*, 106949. https://doi.org/10.1016/j.buildenv.2020.106949",
    ("aries", "2015"): "Aries, M. B. C., Aarts, M. P. J., & van Hoof, J. (2015). Daylight and health: A review of the evidence and consequences for the built environment. *Lighting Research & Technology, 47*(1), 6–27. https://doi.org/10.1177/1477153513509258",
    ("capaldi", "2014"): "Capaldi, C. A., Dopko, R. L., & Zelenski, J. M. (2014). The relationship between nature connectedness and happiness: A meta-analysis. *Frontiers in Psychology, 5*, 976. https://doi.org/10.3389/fpsyg.2014.00976",
    ("carmona", "2010"): "Carmona, M., Gallent, N., & Sarkar, R. (2010). *Space standards: The benefits*. University College London for CABE.",
    ("duarte", "2023"): "Duarte, C. C., Cortiços, N. D., Stefańska, A., & Stefańska, A. (2023). Home balconies during the COVID-19 pandemic: Future architect’s preferences in Lisbon and Warsaw. *Applied Sciences, 13*(1), 298. https://doi.org/10.3390/app13010298",
    ("giang", "2024"): "Giang Thi Ngoc, N., Tsaih, L. S.-J., Chen, J. C.-P., Tamariska, S. R., Coelho, A. M. F., & Kung, H.-Y. (2024). Balcony usage as a space to achieve human well-being during pandemic COVID-19. *Journal of Asian Architecture and Building Engineering*, 1–13. https://doi.org/10.1080/13467581.2024.2370408",
    ("howell", "2013"): "Howell, A. J., & Passmore, H.-A. (2013). The nature of happiness: Nature affiliation and mental well-being. In C. L. M. Keyes (Ed.), *Mental Well-Being* (pp. 231–257). Springer. https://doi.org/10.1007/978-94-007-5195-8_11",
    ("kim", "2021"): "Kim, S., Park, H., & Choo, S. (2021). Effects of changes to architectural elements on human relaxation-arousal responses: Based on VR and EEG. *International Journal of Environmental Research and Public Health, 18*(8), 4305. https://doi.org/10.3390/ijerph18084305",
    ("ko", "2022"): "Ko, W. H., Kent, M. G., Schiavon, S., Levitt, B., & Betti, G. (2022). A window view quality assessment framework. *LEUKOS, 18*(3), 268–293. https://doi.org/10.1080/15502724.2021.1965889",
    ("smektala", "2022"): "Smektała, M., & Baborska-Narożny, M. (2022). The use of apartment balconies: Context, design and social norms. *Buildings and Cities, 3*(1), 134–152. https://doi.org/10.5334/bc.193",
    ("song", "2024"): "Song, T., Xu, L., Zhao, F., & Du, Y. (2024). Healing properties of residential balcony: Characteristics of balcony space in Shanghai’s collective housing. *Journal of Building Engineering, 87*, 108992. https://doi.org/10.1016/j.jobe.2024.108992",
    ("yeom", "2020"): "Yeom, S., Kim, H., Hong, T., Park, H. S., & Lee, D.-E. (2020). An integrated psychological score for occupants according to the windows’ outdoor view size. *Building and Environment, 180*, 107019. https://doi.org/10.1016/j.buildenv.2020.107019",

    
    ("wolkoff", "2018"): "Wolkoff, P. (2018). Indoor air humidity, air quality, and health—An overview. *International Journal of Hygiene and Environmental Health, 221*(3), 376–390. https://doi.org/10.1016/j.ijheh.2018.01.015",
    ("tsai", "2012"): "Tsai, D. H., Lin, J. S., & Chan, C. C. (2012). Office workers’ sick building syndrome and indoor carbon dioxide concentrations. *Journal of Occupational and Environmental Hygiene, 9*(5), 345–351. https://doi.org/10.1080/15459624.2012.673454",
}

def _norm_key_from_citation(cite_text: str):
    
    import re, unicodedata
    m = re.search(r"([A-Za-zÀ-ÖØ-öø-ÿ'\-]+).*?(\d{4})", cite_text)
    if not m:
        return None
    last = m.group(1)
    year = m.group(2)
    
    last = ''.join(c for c in unicodedata.normalize('NFKD', last) if not unicodedata.combining(c)).lower()
    
    if last == "williams" and "cohousing" in cite_text.lower():
        return ("williams-cohousing", year)
    
    last = last.replace("weijs-perrée", "weijs-perree").replace("smektała", "smektala")
    return (last, year)

def extract_citation_keys(text: str):
    
    import re
    keys = set()
    
    for group in re.findall(r"\(([^)]+)\)", text):
        
        for part in re.split(r";|∙|·", group):
            key = _norm_key_from_citation(part)
            if key:
                keys.add(key)
    return keys

def format_references_from_keys(keys, extras=None):
    
    lines = []
    for k in sorted(keys, key=lambda x: (x[0], x[1])):
        if k in MASTER_BIB:
            lines.append(MASTER_BIB[k])
    if extras:
        for s in extras:
            s = s.strip()
            if s and s not in lines:
                lines.append(s)
    return lines


# %%
main_f = Path.home() / "SpinyLeaf_App" 
out_f = main_f / "Wellbeing_Fostered_by_Design"

wb_path = out_f / 'Wellbeing.csv'
com_path = out_f / 'Comfort_Dimension' / 'Comfort.csv'
mat_path = out_f / 'Comfort_Dimension' / 'Materials.csv'
del_path = out_f / 'Delight_Dimension' / 'Delight.csv'
soc_path = out_f / 'Social_Dimension' / 'Social.csv'

wb_im_path = out_f 
com_im_path = out_f / 'Comfort_Dimension'
del_im_path = out_f / 'Delight_Dimension'
soc_im_path = out_f / 'Social_Dimension'

# %%
llm = Llama(
    model_path= "./LLM_Model/Meta-Llama-3-8B-Instruct.Q4_K_M.gguf",
    n_ctx=8192,
    verbose=False
)

def get_completion(prompt):
    response = llm(
        prompt,
        temperature=0.3,
        max_tokens=1024,
        stop=["<|end_of_text|>"],
        repeat_penalty=1.1
    )
    return response['choices'][0]['text'].strip()

# %%
def create_rag_chunks(text):
    """
    Splits a structured reference text into RAG-friendly chunks.
    Each chunk is a paragraph or bullet group.
    """
    raw_chunks = text.split("\n\n")
    chunks = [chunk.strip() for chunk in raw_chunks if chunk.strip()]
    return chunks

# %%
def format_llama3_prompt(user_prompt):
    return (
        "<|begin_of_text|>"
        "<|start_header_id|>user<|end_header_id|>\n"
        f"{user_prompt}\n"
        "<|start_header_id|>assistant<|end_header_id|>\n"
    )

# %%
def df_to_dict(path, cols=None, filter_col=None, filter_val=None, exclude_prefixes=None):
    df = pd.read_csv(path)

    if cols:
        df = df[cols]

    if filter_col and filter_val is not None:
        df = df[df[filter_col] < filter_val]

    if exclude_prefixes:
        pattern = '|'.join(f"^{prefix}" for prefix in exclude_prefixes)
        df = df[~df["room_ids"].str.contains(pattern, regex=True)]

    return df.to_dict(orient="records")

# %%
def extract_wellbeing_summary(path):
    df = pd.read_csv(path)

    avg = round(df["wellbeing_satisfaction"].mean(), 2)

    satisfied = df[df["wellbeing_satisfaction"] >= 4].shape[0]
    neutral = df[(df["wellbeing_satisfaction"] >= 2) & (df["wellbeing_satisfaction"] < 4)].shape[0]
    dissatisfied = df[df["wellbeing_satisfaction"] < 2].shape[0]

    low_rooms_df = df[df["wellbeing_satisfaction"] < 3][["room_ids", "wellbeing_satisfaction"]]
    low_rooms = ', '.join(
        f'{row.room_ids} ({round(row.wellbeing_satisfaction, 2)})'
        for _, row in low_rooms_df.iterrows()
    ) or "None"

    mean_comfort = round(df["comfort_satisfaction"].mean(), 2)
    mean_delight = round(df["delight_satisfaction"].mean(), 2)
    mean_social = round(df["social_satisfaction"].mean(), 2)

    return avg, satisfied, neutral, dissatisfied, low_rooms, mean_comfort, mean_delight, mean_social

# %%
def extract_comfort_summary(path):
    df = pd.read_csv(path)

    means = {
        "thermal": round(df["extreme_hot_satisf"].mean(), 2),
        "daylight": round(df["daylight_satisf"].mean(), 2),
        "acoustic": round(df["sound_satisf"].mean(), 2),
        "air_quality": round(df["air_quali"].mean(), 2)
    }

    issues = []
    for factor, col in {
        "thermal": ["extreme_hot_satisf"],
        "daylight": ["daylight_satisf"],
        "acoustic": ["sound_satisf"],
        "air_quality": ["air_quali"]
    }.items():
        for c in col:
            df_issue = df[df[c] < 0.66][["room_ids", "floor_level", c]]
            for _, row in df_issue.iterrows():
                issues.append(f"- {row.room_ids} (Floor {row.floor_level}): {factor} = {round(row[c], 2)}")

    factor_issue_counts = {
        "thermal": (df["extreme_hot_satisf"] < 0.66).sum(),
        "daylight": (df["daylight_satisf"] < 0.66).sum(),
        "acoustic": (df["sound_satisf"] < 0.66).sum(),
        "air_quality": (df["air_quali"] < 0.66).sum()
    }

    max_issues = max(factor_issue_counts.values())
    worst_factors = [factor for factor, count in factor_issue_counts.items() if count == max_issues and count > 0]

    return means, issues, factor_issue_counts, worst_factors

# %%
def extract_delight_summary(path):
    df = pd.read_csv(path)

    means = {
        "views": round(df["views_overall_satisf"].mean(), 2),
        "balcony": round(df["balcony_satisf"].mean(), 2),
        "space_size": round(df["space_size_satisf"].mean(), 2)
    }

    issues = []
    factor_columns = {
        "views": "views_overall_satisf",
        "balcony": "balcony_satisf",
        "space_size": "space_size_satisf"
    }

    for factor, column in factor_columns.items():
        df_issue = df[(df[column] < 0.66) & (~df["room_ids"].str.startswith("CORE"))][["room_ids", "floor_level", column]]
        for _, row in df_issue.iterrows():
            issues.append(f"- {row.room_ids} (Floor {row.floor_level}): {factor} = {round(row[column], 2)}")

    factor_issue_counts = {
        factor: ((df[column] < 0.66) & (~df["room_ids"].str.startswith("CORE"))).sum()
        for factor, column in factor_columns.items()
    }

    max_issues = max(factor_issue_counts.values())
    worst_factors = [factor for factor, count in factor_issue_counts.items() if count == max_issues and count > 0]

    return means, issues, factor_issue_counts, worst_factors

# %%
def extract_social_summary(path):
    df = pd.read_csv(path)
    means = {
        "social_amount": round(df["social_amount_satisf"].mean(), 2),
        "social_distribution": round(df["social_distribution_satisf"].mean(), 2),
        "social_green": round(df["social_green_satisf"].mean(), 2)
    }

    issues = []
    factor_columns = {
        "social_amount": "social_amount_satisf",
        "social_distribution": "social_distribution_satisf",
        "social_green": "social_green_satisf"
    }

    for factor, column in factor_columns.items():
        df_issue = df[df[column] < 0.66][["room_ids", "floor_level", column]]
        for _, row in df_issue.iterrows():
            issues.append(f"- {row.room_ids} (Floor {row.floor_level}): {factor} = {round(row[column], 2)}")

    factor_issue_counts = {
        factor: (df[column] < 0.66).sum()
        for factor, column in factor_columns.items()
    }

    max_issues = max(factor_issue_counts.values())
    worst_factors = [factor for factor, count in factor_issue_counts.items() if count == max_issues and count > 0]

    return means, issues, factor_issue_counts, worst_factors

# %%
def extract_materials_summary_enhanced(path):
    df = pd.read_csv(path)
    row = df.iloc[0]

    window_type_str = row["Window_type"]
    glazing_match = re.search(r'(Sgl|Dbl|Trp)', window_type_str)
    glazing = {
        'Sgl': 'Single glazing',
        'Dbl': 'Double glazing',
        'Trp': 'Triple glazing'
    }.get(glazing_match.group(1), 'Unknown glazing') if glazing_match else 'Unknown glazing'

    glass_type = window_type_str.split(glazing_match.group(1))[1].strip() if glazing_match else window_type_str

    summary = {
        "glazing_type": glazing,
        "glass_type": glass_type,
        "windows_u": round(row["Windows_U"], 2),
        "window_noise_reduction": round(row["Win_reduction"], 1),
        "shgc": round(row["SHGC"], 2),
        "wall_r_insulation": row["Wall_R"],
        "wall_noise_reduction": row["Wall_reducrion"],
        "roof_r_insulation": row["Roof_R"],
        "ground_r_insulation": row["Ground_R"]
    }

    return summary

# %%
exclude = ["CORE", "SOCIAL"]

wellbeing_df = df_to_dict(
    wb_path,
    cols=["room_ids", "floor_area", "wellbeing_satisfaction", "comfort_satisfaction", "delight_satisfaction", "social_satisfaction"],
    exclude_prefixes=exclude
)

comfort_df = df_to_dict(
    com_path,
    cols=["room_ids", "floor_level", "extreme_hot_satisf", "extreme_cold_satisf", "daylight_satisf", "sound_satisf", "air_quali", "comfort_satisfaction"],
    exclude_prefixes=exclude
)

delight_df = df_to_dict(
    del_path,
    cols=["room_ids", "floor_level", "hor_views_satisf", "sky_views_satisf", "green_views_satisf", "balcony_satisf", "space_size_satisf", "delight_satisfaction"],
    exclude_prefixes=exclude
)

social_df = df_to_dict(
    soc_path,
    cols=["room_ids", "floor_level", "social_amount_satisf", "social_distribution_satisf", "social_green_satisf", "social_satisfaction"],
    exclude_prefixes=exclude
)

# %%
thermal_reference = f"""
1. Thermal Comfort Standards

Thermal Comfort – Research Data Analysis (Numerical) showed that:

Apartments (Residents): Satisfaction most likely between 19.3 °C and 28.3 °C. Dissatisfaction most common below 16 °C or above 31.6 °C.

Offices (Workers): Satisfaction most likely between 20.8 °C and 26.2 °C. Dissatisfaction increases below 18.5 °C and above 28.3 °C.

Thermal Comfort – Research Data Analysis (Generalised)

Apartment residents reported satisfaction across a wider thermal range than office workers.

Residents are generally comfortable between cooler and warmer temperatures, while office workers indicated a narrower comfort band, preferring conditions closer to the mid-20s °C.

Both groups expressed dissatisfaction at lower extremes (below ~18 °C) and higher extremes (above ~30 °C).

2. Window Performance
Glazing Types:

Double Glazing: Two panes of glass with an air or gas-filled space between them, reducing heat transfer.

Triple Glazing: Adds a third pane, further enhancing insulation and energy efficiency.

U-Value:

The U-value measures the rate of heat transfer; lower values indicate better insulation.

Double Glazing: Typically has U-values ranging from 1.2 to 3.7 W/m²·K.

Triple Glazing: Often achieves U-values below 1.0 W/m²·K, offering superior thermal performance.

Low-Emissivity (Low-E) Coatings:

Low-E coatings are microscopically thin layers applied to glazing surfaces to reduce infrared radiation, minimizing heat transfer without compromising visible light transmission.

Solar Heat Gain Coefficient (SHGC):

SHGC measures how much solar radiation passes through a window.

Hot Climates: Lower SHGC (<0.35) reduces cooling loads.

Cold Climates: Higher SHGC (0.45–0.55) allows for passive solar heating.

3. Insulation Strategies
Wall Insulation:

Proper wall insulation is crucial for thermal comfort.

R-Values: For temperate climates, R-values above R-2.5–R-3.5 (m²·K/W) are considered minimal, with higher values recommended for colder regions.

Roof Insulation:

Roof assemblies significantly impact thermal buffering.

R-Values: Roof insulation should reach R-4.0 or higher in residential buildings, especially in top-floor dwellings.
"""

# %%
daylight_reference = f"""
Daylighting Design Guidelines for Enhanced Indoor Comfort and Energy Efficiency
Daylighting is the strategic use of natural light to illuminate building interiors. Its goal is to improve occupant well-being and reduce energy consumption by minimizing reliance on electric lighting. Effective daylighting design considers building orientation, glazing properties, window placement, interior reflectance, and glare control.

Standards and Guidelines
LEED v4 (Leadership in Energy and Environmental Design) awards credits for daylight access, visual comfort, and quality views.

Design Strategies
Window-to-Wall Ratio (WWR):

Maintain a balanced WWR to allow daylight penetration without excessive heat gain or glare.

A WWR between 20% and 40% is often recommended depending on orientation and climate zone.

Glazing Selection:

Use high-performance glazing with Visible Transmittance (VT) of 0.6–0.7 for good daylight entry.

Select low SHGC (Solar Heat Gain Coefficient) glass to minimize overheating in warmer climates.

Shading Devices:

Integrate fixed or dynamic shading systems such as overhangs, light shelves, and louvers.

These reduce direct glare and redirect light deeper into interior spaces.

Interior Reflectance:

Choose light-colored wall and ceiling finishes to improve light diffusion.

Keep partitions low or translucent to extend daylight access across open-plan layouts.
"""

# %%
acoustic_reference = f"""
Acoustic Design Guidelines for Building Performance
1. Acoustic Performance Standards
Effective acoustic design in buildings is guided by several key standards:

ANSI/ASA S12.60: Establishes acoustic performance criteria for classrooms, recommending maximum background noise levels of 35 dBA.

LEED and WELL Building Standards: Incorporate acoustic performance as a criterion for certification, promoting environments that support occupant comfort and wellbeing.

Recommended Noise Limits
International standards and guidelines define recommended indoor noise levels for comfort and health:

WHO Environmental Noise Guidelines (2018):

≤ 35 dBA in day-use rooms (e.g., offices, living rooms)

≤ 30 dBA in bedrooms at night

Outdoor façade exposure should not exceed 53 dBA Lden in urban environments (preferably <50 dBA)

ASHRAE Handbook – HVAC Applications (2023):

Recommends Noise Criteria (NC) levels of 25–35 for residential and office spaces

ISO 10052 / ISO 16283-3 (field measurement):

Establish procedures for evaluating outdoor-to-indoor sound level differences (DnT,w or D2m,nT,w)

Key Design Strategies 
1. Glazing Selection

Avoid single glazing in noise-exposed areas

Instead, use:

Double glazing (≥6/12/6 mm): STC 32–36
Laminated glass: significantly reduces higher-frequency noise
Asymmetric glazing (e.g., 6mm + 10mm): better at controlling low-frequency traffic noise

Laminated Glass: Incorporate laminated glazing, which includes a viscoelastic interlayer, to enhance sound insulation properties.

Double or Triple Glazing: Use multiple glazing layers with air or inert gas fills to improve acoustic performance.

Window-to-Wall Ratio: Optimize the size and placement of windows to balance natural light and acoustic insulation.

2. Wall Systems (Indicative R Values)
If walls are known to be lightweight (e.g., framed construction):
Add mass or specify external cladding with high Rw (≥45 dB)
Use staggered stud designs or double-skin façades (if known early)
Mineral Wool and Fiberglass Insulation: Effective for absorbing sound within wall cavities.

3. Openings and Penetrations
Limit use of fixed louvres or vents on exposed façades unless acoustically treated
Ensure window-to-wall ratio is minimized on noisy façades
"""

# %%
airquali_reference = f"""
Indoor relative humidity (RH) plays a critical role in shaping indoor air quality (IAQ) and associated health outcomes. A consistent body of literature indicates that maintaining RH between 40% and 60% is generally optimal for occupant comfort and health (Wolkoff, 2018). 
There is also evidence suggesting a link between moderate CO₂ elevations and Sick Building Syndrome symptoms (Tsai et al., 2012). 
Maintaining levels below 1000 ppm is widely regarded as best practice to prevent fatigue, mitigate symptoms of Sick Building Syndrome, and support cognitive performance. 
Strategies such as demand-controlled ventilation and adequate outdoor air supply are essential to manage indoor CO₂ levels effectively.

Enable cross-ventilation through operable windows on opposing façades.
Size and locate openings to capture prevailing winds (especially in residential and mixed-mode spaces)

In spaces with non-operable windows, special attention must be given to ensuring adequate mechanical ventilation.
Provide space allowances for future mechanical ventilation, even in naturally ventilated schemes

3. Prevent RH Extremes through Envelope & Massing
Incorporate thermal mass and moderate glazing ratios to stabilise temperature and reduce condensation
Avoid unshaded large glazed surfaces that can drive up RH from condensation or limit fresh air access due to overheating concerns
Locate bathrooms/kitchens along external walls where moisture can be vented directly outdoors

References:

Wolkoff, P., 2018. Indoor air humidity, air quality, and health–An overview. *International Journal of Hygiene and Environmental Health, 221*(3), 376–390.

Tsai, D. H., Lin, J. S., & Chan, C. C., 2012. Office workers’ sick building syndrome and indoor carbon dioxide concentrations. *Journal of Occupational and Environmental Hygiene, 9*(5), 345–351.
"""

# %%
delight_reference = f"""
Views of the outdoors:
Attractive outdoor views, particularly of the sky and natural scenes, support circadian wellbeing, visual comfort, cognitive performance, and overall satisfaction (Altomonte et al., 2020; Aries et al., 2015). 
Occupants generally prefer seating near windows, with higher satisfaction reported when window-to-wall ratios exceed 25% (Ko et al., 2022; Yeom et al., 2020). 
Larger ratios above ~65% bring limited additional benefit (Kim et al., 2021). 
Real-world experiences highlight that long-distance and dynamic views are important for daily wellbeing, which should be considered in early design through orientation, floorplate depth, and window placement.

Direct contact with nature:
Exposure to nature provides restorative effects, reduces stress, and improves wellbeing across diverse contexts (Capaldi et al., 2014; Howell & Passmore, 2013). 
In dense cities, private access can be enhanced through balconies, rooftop gardens, and vertical greenery systems. 
Balconies gained new significance during the pandemic, supporting residents’ wellbeing (Duarte et al., 2023; Giang Thi Ngoc et al., 2024). 
Their usability depends strongly on size, with spaces above ~10–15 m² used more frequently and perceived as more supportive (Smektała & Baborska-Narożny, 2022; Song et al., 2024).

Living-space size:
Adequate dwelling size is linked to family dynamics, privacy, social life, and educational performance (Carmona et al., 2010). 
Overcrowding can negatively affect children’s learning and household relationships. In high-rise developments, minimised floor areas increase these risks, making early design decisions on layout, depth, and ceiling height critical.
"""

# %%
social_reference = f"""
Opportunities for social interactions:
Social connections are fundamental to wellbeing, with strong ties and supportive networks linked to better mental and physical health (Jordan, 2023; Umberson & Montez, 2010). 
Loneliness and isolation are major health risks, and even minimal interactions or community-based initiatives can enhance wellbeing (Gunaydin et al., 2021; Leavell et al., 2019).
The built environment plays a central role: walkable neighbourhoods and accessible social spaces foster interaction (Weijs-Perrée et al., 2015), while poorly designed high-rise housing has been associated with isolation, stress, and reduced cohesion (Barros et al., 2019; Kearns et al., 2012). 
Recent studies highlight how older adults, students, and low-income residents are particularly vulnerable when social infrastructure is lacking (Nguyen et al., 2025, 2024).
Design strategies such as transitional or semi-private zones, vertical social pockets, and circulation areas can mitigate social overload while encouraging connection (Bee & Im, 2016; Williams, 2005). 
The quality, size, and connectivity of communal areas are critical for fostering belonging, neighbourliness, and positive emotions (Barrie et al., 2023; Kleeman et al., 2023). 
Successful public and shared spaces work best when well-designed and actively supported by community programming, helping to strengthen cohesion and reduce loneliness in high-rise living.
"""

# %%
def prompt_wellbeing_2(avg, satisfied, neutral, dissatisfied, low_rooms,
                       mean_comfort, mean_delight, mean_social,
                       ranking_text, lowest_dimension, lowest_score):
    prompt = f"""

You are an expert in building performance and wellbeing evaluation.

Below is a summary of spatial wellbeing data from a building. All scores are normalised:

- The overall wellbeing score ranges from **0 to 6**
- mean wellbeing < 2: dissatisfied
- mean wellbeing >= 2 and < 4: neutral
- mean wellbeing >= 4: satisfied

- The satisfaction scores for each dimension (comfort, delight, social) range from **0 to 2**
- < 0.66 = dissatisfied
- 0.66 to <1.33 = neutral
- ≥ 1.33 = satisfied

---

## Wellbeing Fostered By Design Summary
- Explaind that the Wellbeing score ranges from 0 to 6
- Mean wellbeing score: {avg}
- Room ratings: {satisfied} satisfied, {neutral} neutral, {dissatisfied} dissatisfied
- Problematic rooms: {low_rooms}
- Explaind that the Comfort, Delight and Social scores ranges from 0 to 2
- Comfort mean: {mean_comfort}
- Delight mean: {mean_delight}
- Social mean: {mean_social}

---

Your Task is to write a formal report that should include the following:

- Building Wellbeing Fostered by Design summary  
- Critical rooms and issues  
- Rank of dimensions from weakest to strongest:  
{ranking_text}

- The last paragraph must point out which dimension most needs attention:  
The dimension with the lowest mean is **{lowest_dimension}** ({lowest_score}).

Add the placeholder: `wellbeing_images`

---

Output Rules
- Use a formal, professional tone
- Output the report only with no introductions or follow up messages
- Ensure you explain that the Wellbeing score scale range from 0 to 6
- Ensure you explain that the Comfort, Delight and Social scores ranges from 0 to 2
- Ensure the image placeholder is added only once and is written exactly as: `wellbeing_images`
"""
    return prompt

# %%
def prompt_comfort_1(means, issues, factor_issue_counts, worst_factors, materials):
    if isinstance(worst_factors, list):
        worst_factors_text = ", ".join([f"**{f}**" for f in worst_factors])
        worst_counts_text = ", ".join([f"{factor_issue_counts[f]} for {f}" for f in worst_factors])
    else:
        worst_factors_text = f"**{worst_factors}**"
        worst_counts_text = f"{factor_issue_counts[worst_factors]}"

    prompt = f"""
You are an expert in indoor environmental quality and comfort design.

Scores for comfort factors (thermal, daylight, acoustic, air quality) range from 0 to 2:
- < 0.66 = dissatisfied
- 0.66 to <1.33 = neutral
- ≥ 1.33 = satisfied

---

## Materials Summary
Construction and glazing characteristics:
- Glazing: {materials['glazing_type']}
- Glass type: {materials['glass_type']}
- Window U-value: {materials['windows_u']}
- SHGC: {materials['shgc']}
- Window noise reduction: {materials['window_noise_reduction']} dB
- Wall insulation R-value: {materials['wall_r_insulation']}
- Wall noise reduction: {materials['wall_noise_reduction']} dB
- Roof insulation R-value: {materials['roof_r_insulation']}
- Ground insulation R-value: {materials['ground_r_insulation']}

## Comfort Dimension
Mean Comfort Scores:
- Thermal: {means['thermal']}
- Daylight: {means['daylight']}
- Acoustic: {means['acoustic']}
- Air Quality: {means['air_quality']}

The lowest-performing factor(s): {worst_factors_text}, with {worst_counts_text} rooms rated as dissatisfied.

---

Write a Comfort Dimension Summary:
- Start with ## Comfort Dimension Summary
- Describe the overall Comfort performance
- Emphasise which factors performed worst and why
- Refer to the 0–2 scoring scale for interpretation
- Write the Materials Summary as provided
- Add the single placeholder at the end: `comfort_images`

---

Output Rules
- Use formal, professional tone
- Output the report only with no introductions or follow up messages
- Don't write in first person
- Do not use email or message style
"""
    return prompt

# %%
def prompt_thermal(thermal_issues, materials, thermal_chunks):
    issue_text = ""
    recommendation_text = ""

    if thermal_issues:
        issue_text = (
            "The following rooms showed dissatisfaction (score < 0.66) in Thermal Comfort:\n"
            + "\n".join(thermal_issues)
        )

        recommendation_text = f"""
## Recommendations for Improving Thermal Comfort

Write clear and actionable recommendations for designers to improve thermal comfort based on the issues above. Justify your suggestions using the knowledge section above.

- The recommendation should not include furniture, layout or finishing materials.  
- Do not recommend retrofit strategies, it is not an existent building.
- All rooms listed above are located in the same building and use the same construction materials. Therefore, write general building-wide recommendations to improve thermal comfort, without repeating room-specific instructions.
- When recommending upgrades, consider the current materials:
  • If the glazing is single, suggest upgrading to **double glazing with Low-E coating**. Only recommend **triple glazing** for extreme conditions.
  • Use SHGC and U-value thresholds from the materials summary to determine if improvements are needed.
  • If wall or roof insulation values are already moderate (e.g., R ≥ 3), only recommend improvement if thermal comfort is still low.
  • Suggest **non-material strategies** (e.g., shading, ventilation, thermal zoning).
"""
    else:
        issue_text = "All rooms are performing well in terms of thermal comfort. No improvements are currently required."

    prompt = f"""
You are an expert in building performance and indoor thermal comfort.

Thermal satisfaction is scored from 0 to 2:
- Below 0.66: dissatisfied
- 0.66 to <1.33: neutral
- 1.33 or higher: satisfied

---

## Supporting Knowledge

Use the following expert guidance and literature references to justify your recommendations:

{chr(10).join(thermal_chunks)}

---
Start the report as:

## Thermal Comfort Detected Issues
{issue_text}

---

Construction and glazing characteristics:
- Glazing: {materials['glazing_type']}
- Glass type: {materials['glass_type']}
- Window U-value: {materials['windows_u']}
- SHGC: {materials['shgc']}
- Wall insulation R-value: {materials['wall_r_insulation']}
- Roof insulation R-value: {materials['roof_r_insulation']}
- Ground insulation R-value: {materials['ground_r_insulation']}

{recommendation_text}

---

Output Rules
- Use formal, professional tone
- Output the report only with no introductions or follow up messages
- Do not write in first person
- Do not use email or message style
"""
    return prompt

# %%
def prompt_daylight(daylight_issues, materials, daylight_chunks):
    issue_text = ""
    recommendation_text = ""

    if daylight_issues:
        issue_text = (
            "The following rooms showed dissatisfaction (score < 0.66) in Daylight Access:\n"
            + "\n".join(daylight_issues)
        )

        recommendation_text = """
## Recommendations for Improving Daylight Access

Write clear and actionable recommendations for designers to improve daylight access based on the issues above. Justify your suggestions using the knowledge section above.

- The recommendations should not include furniture, layout, or finishing materials.
- Do not recommend retrofit strategies, it is not an existent building.
- All rooms listed above are located in the same building and use the same construction materials. Therefore, write a general building-wide recommendation to improve daylight access, without repeating room-specific instructions.
- Write the recommendations according to the references.

Conclude after listing recommendations.
"""
    else:
        issue_text = "All rooms are performing well in terms of daylight access. No improvements are currently required."

    prompt = f"""
You are an expert in daylighting and visual comfort design.

Comfort scores for daylight range from 0 to 2:
- Below 0.66: dissatisfied
- 0.66 to <1.33: neutral
- 1.33 or higher: satisfied

---

## Supporting Knowledge

Use the following expert guidance and best-practice references to inform your recommendations:
{chr(10).join(daylight_chunks)}

---
Start the report as:

## Daylight Access Detected Issues
{issue_text}

---

- Glazing: {materials['glazing_type']}
- Glass type: {materials['glass_type']}
- Visible Transmittance (approx): {materials.get('vt', 'unknown')}
- SHGC: {materials['shgc']}

{recommendation_text}

---

Output Rules
- Use formal, professional tone
- Output the report only with no introductions or follow up messages
- Do not write in first person
- Do not use email or message style
"""
    return prompt

# %%
def prompt_acoustic(acoustic_issues, materials, acoustic_chunks):
    issue_text = ""
    recommendation_text = ""

    if acoustic_issues:
        issue_text = (
            "The following rooms showed dissatisfaction (score < 0.66) in Acoustic Comfort:\n"
            + "\n".join(acoustic_issues)
        )

        recommendation_text = """

## Recommendations for Improving Acoustic Comfort

Write clear and actionable recommendations for designers to improve acoustic comfort based on the issues above. Justify your suggestions using the knowledge section above.

- The recommendation should not include furniture, layout or finishing materials.
- Do not recommend retrofit strategies, it is not an existent building.
- All rooms listed above are located in the same building and use the same construction materials. Therefore, write a general building-wide recommendation to improve acoustic comfort, without repeating room-specific instructions.
- Write the recommendations according to the references.

Conclude after listing recommendations.
"""
    else:
        issue_text = "All rooms are performing well in terms of acoustic comfort. No improvements are currently required."

    prompt = f"""
You are an expert in acoustic comfort design.

Comfort scores for acoustic range from 0 to 2:
- Below 0.66: dissatisfied
- 0.66 to <1.33: neutral
- 1.33 or higher: satisfied

---

## Supporting Knowledge

Use the following expert guidance and best-practice references to inform your recommendations:
{chr(10).join(acoustic_chunks)}

---
Start the report as:

## Acoustic Comfort Detected Issues
{issue_text}

---

- Window noise reduction: {materials['window_noise_reduction']} dB
- Wall noise reduction: {materials['wall_noise_reduction']} dB

{recommendation_text}

---

Output Rules
- Use formal, professional tone
- Output the report only with no introductions or follow up messages
- Do not write in first person
- Do not use email or message style
"""
    return prompt

# %%
def prompt_airquali(airquali_issues, materials, airquali_chunks):
    issue_text = ""
    recommendation_text = ""

    if airquali_issues:
        issue_text = (
            "The following rooms showed dissatisfaction (score < 0.66) in Air Quality:\n"
            + "\n".join(airquali_issues)
        )

        recommendation_text = """
## Recommendations for Improving Air Quality

Write clear and actionable recommendations for designers to improve Air Quality based on the issues above. Justify your suggestions using the knowledge section above.

- The recommendation should not include furniture, layout or finishing materials. 
- Do not recommend retrofit strategies, it is not an existent building.    
- All rooms listed above are located in the same building and use the same construction materials. Therefore, write general building-wide recommendations to improve air quality, without repeating room-specific instructions.
- Write the recommendations according to the references

Conclude after listing recommendations.
"""
    else:
        issue_text = "All rooms are performing well in terms of air quality. No improvements are currently required."

    prompt = f"""
You are an expert in air quality and indoor environmental design.

Comfort scores for air quality range from 0 to 2:
- Below 0.66: dissatisfied
- 0.66 to <1.33: neutral
- 1.33 or higher: satisfied

---

## Supporting Knowledge

Use the following expert guidance and best-practice references to inform your recommendations:
{chr(10).join(airquali_chunks)}

---
Start the report as:

## Air Quality Detected Issues
{issue_text}

{recommendation_text}

---

Output Rules
- Use formal, professional tone
- Output the report only with no introductions or follow up messages
- Do not write in first person
- Do not use email or message style
"""
    return prompt

# %%
def prompt_delight(means, issues, factor_issue_counts, worst_factors, delight_chunks):
    grouped_issues = {
        "views": [],
        "balcony": [],
        "space_size": []
    }

    for line in issues:
        for factor in grouped_issues:
            if factor in line:
                grouped_issues[factor].append(line)

    if isinstance(worst_factors, list):
        worst_factors_text = ", ".join([f"**{f}**" for f in worst_factors])
    else:
        worst_factors_text = f"**{worst_factors}**"

    prompt = f"""
You are an expert in building performance and positive stimuli.

Evaluate the building performance regarding the Delight dimension:
Delight: views, access to green space / balconies, and space size

Delight scores range from 0 to 2:
- Below 0.66: dissatisfied
- 0.66 to <1.33: neutral
- 1.33 or higher: satisfied

---

## Supporting Knowledge

Use the following expert guidance and best-practice references to inform your recommendations:
{chr(10).join(delight_chunks)}
- You must cite the relevant sources from the supporting knowledge when appropriate.
- Use APA-style in-text citations, e.g., (Kaplan, 1995), when referring to supporting studies or guidance.
- The complete References must be cited at the end in APA style.
- Do not invent or cite references that are not present in the supporting knowledge.
- All recommendations must be justified by one or more sources from the supporting material.
- Avoid generic advice; tie each suggestion to specific references when applicable.

---
Start the report as:

## Delight Dimension

Mean delight scores:
- Views: {means['views']}
- Direct Access to Green spaces (Balconies): {means['balcony']}
- Space Size: {means['space_size']}

The lowest-performing factor(s): {worst_factors_text}

---

Insert `delight_images` at the end of this section.

---

## Detected Issues and Recommendations
"""

    for factor, label, recommendation in [
        ("views", "Views", "increase number or size of windows; if in dense contexts, consider adding balconies or increasing room size to mitigate"),
        ("balcony", "Direct Access to Green spaces (Balconies)", "increase balcony size or provide direct access to private/shared outdoor spaces"),
        ("space_size", "Space Size", "increase apartment size for (RESID) rooms or reduce occupancy density for (COMMERC) rooms")
    ]:
        entries = grouped_issues[factor]
        if entries:
            prompt += f"""

### {label}
Rooms with dissatisfaction:
{chr(10).join(entries)}

Recommendations:
- {recommendation}. Support your recommendation using references from the knowledge section above. Use APA-style in-text citations.

add the complete references for all the citations in the recommendations at the end of the report, the section must start with ## References
"""
        else:
            prompt += f"""

### {label}
All rooms are performing well in this aspect. No improvements required.
"""

    prompt += """

---

Output Rules:
- Use formal, professional tone
- Output the report only with no introductions or follow up messages
- Do not write in first person
- Do not use email or message style
- Ensure image placeholder is exactly: `delight_images`
"""
    return prompt

# %%
def prompt_social(means, issues, factor_issue_counts, worst_factors, social_chunks):
    grouped_issues = {
        "social_amount": [],
        "social_distribution": [],
        "social_green": []
    }

    for line in issues:
        if "social_amount" in line:
            grouped_issues["social_amount"].append(line)
        elif "social_distribution" in line:
            grouped_issues["social_distribution"].append(line)
        elif "social_green" in line:
            grouped_issues["social_green"].append(line)

    if isinstance(worst_factors, list):
        worst_factors_text = ", ".join([f"**{wf}**" for wf in worst_factors])
    else:
        worst_factors_text = f"**{worst_factors}**"

    prompt = f"""
You are an expert in building performance and social space design.

Evaluate the building performance regarding the Social dimension:
Social: Amount of Social Spaces, Social Green Spaces/ Gardens, Distribution of Social Spaces

Social scores range from 0 to 2:
- Below 0.66: dissatisfied
- 0.66 to <1.33: neutral
- 1.33 or higher: satisfied

---

## Supporting Knowledge

Use the following expert guidance and best-practice references to inform your recommendations:
{chr(10).join(social_chunks)}
- You must cite the relevant sources from the supporting knowledge when appropriate.
- Use APA-style in-text citations, e.g., (Kaplan, 1995), when referring to supporting studies or guidance.
- The complete References must be cited at the end in APA style.
- Do not invent or cite references that are not present in the supporting knowledge.
- All recommendations must be justified by one or more sources from the supporting material.
- Avoid generic advice; tie each suggestion to specific references when applicable.

---
Start the report as:

## Social Dimension

Mean social scores:
- Amount of Social Spaces: {means['social_amount']}
- Social Green Spaces/ Gardens: {means['social_green']}
- Distribution of Social Spaces: {means['social_distribution']}

The lowest-performing factor(s): {worst_factors_text}

Insert `social_images` at the end of this section.

---

## Detected Issues and Recommendations
"""

    for factor, label, recommendation in [
        ("social_amount", "Amount of Social Spaces", "increase indoor/outdoor social space areas"),
        ("social_distribution", "Distribution of Social Spaces", "add transitional areas (e.g., SOCIAL_L3 or SOCIAL_L4)"),
        ("social_green", "Social Green Spaces/ Gardens", "add communal green spaces (e.g., rooftops)")
    ]:
        entries = grouped_issues[factor]
        if entries:
            prompt += f"""

### {label}
Rooms with dissatisfaction:
{chr(10).join(entries)}

Recommendations:
- {recommendation}. Support your recommendation using references from the knowledge section above. Use APA-style in-text citations.

add the complete references for all the citations in the recommendations at the end of the report, the section must start with ## References
"""
        else:
            prompt += f"""

### {label}
All rooms are performing well in this aspect. No improvements required.
"""

    prompt += """

---

Output Rules
- Use formal, professional tone
- Do not write in first person
- Output the report only with no introductions or follow up messages
- Do not use email or message style
- Ensure image placeholder is exactly: `social_images`
"""

    return prompt

# %%
avg, satisfied, neutral, dissatisfied, low_rooms, mean_comfort, mean_delight, mean_social = extract_wellbeing_summary(wb_path)

# %%
dimensions = {
    "Comfort": mean_comfort,
    "Delight": mean_delight,
    "Social": mean_social
}
sorted_dims = sorted(dimensions.items(), key=lambda x: x[1])
ranking_text = '\n'.join([f"{i+1}. {dim} ({score})" for i, (dim, score) in enumerate(sorted_dims)])
lowest_dimension = sorted_dims[0][0]
lowest_score = sorted_dims[0][1]

# %%
c_means, c_issues, c_factor_issue_counts, c_worst_factor = extract_comfort_summary(com_path)

# %%
thermal_issues = [issue for issue in c_issues if "thermal" in issue.lower()]
daylight_issues = [issue for issue in c_issues if "daylight" in issue.lower()]
acoustic_issues = [issue for issue in c_issues if "acoustic" in issue.lower()]
airquali_issues = [issue for issue in c_issues if "air_quality" in issue.lower()]

# %%
materials = extract_materials_summary_enhanced(mat_path)

# %%
wellbeing_text = prompt_wellbeing_2(avg, satisfied, neutral, dissatisfied, low_rooms,
                       mean_comfort, mean_delight, mean_social,
                       ranking_text, lowest_dimension, lowest_score)
wellbeing_text_formatted = format_llama3_prompt(wellbeing_text)
wellbeing_report = get_completion(wellbeing_text_formatted)

# %%
comfort_text_1 = prompt_comfort_1(c_means, c_issues, c_factor_issue_counts, c_worst_factor, materials)
comfort_text_formatted = format_llama3_prompt(comfort_text_1)
comfort_intro = get_completion(comfort_text_formatted)

chk = 2

# %%
thermal_rag_chunks = create_rag_chunks(thermal_reference)
keywords = ["thermal", "ashrae", "temperature", "u-value", "shgc", "glazing", "low-e", "insulation", "r-value", "roof", "wall", "heat gain", "heat loss"]
thermal_chunks = [chunk for chunk in thermal_rag_chunks if any(kw in chunk.lower() for kw in keywords)]
thermal_text = prompt_thermal(thermal_issues, materials, thermal_chunks[:chk])
thermal_text_formatted = format_llama3_prompt(thermal_text)
thermal_report = get_completion(thermal_text_formatted)

# %%
dl_rag_chunks = create_rag_chunks(daylight_reference)
keywords = ["daylight", "wwr", "window", "visible transmittance", "vt", "glare", "shading", "light shelf", "louver", "overhang", "reflectance", "ashrae", "leed"]
daylight_chunks = [chunk for chunk in dl_rag_chunks if any(kw in chunk.lower() for kw in keywords)]
daylight_text = prompt_daylight(daylight_issues, materials, daylight_chunks[:chk])
daylight_text_formatted = format_llama3_prompt(daylight_text)
daylight_report = get_completion(daylight_text_formatted)

# %%
acoustic_rag_chunks = create_rag_chunks(acoustic_reference)
keywords = [ "acoustic", "noise", "sound", "db", "barrier", "soundproof", "insulation", "glazing", "sealed window"]
acoustic_chunks = [chunk for chunk in acoustic_rag_chunks if any(kw in chunk.lower() for kw in keywords)]
acoustic_text = prompt_acoustic(acoustic_issues, materials, acoustic_chunks[:chk])
acoustic_text_formatted = format_llama3_prompt(acoustic_text)
acoustic_report = get_completion(acoustic_text_formatted)

# %%
aq_rag_chunks = create_rag_chunks(airquali_reference)
keywords = ["air quality", "co2", "ventilation", "natural ventilation", "mechanical ventilation", "humidity", "fresh air", "ashrae", "well"]
air_quality_chunks = [chunk for chunk in aq_rag_chunks if any(kw in chunk.lower() for kw in keywords)]
airquali_text = prompt_airquali(airquali_issues, materials, air_quality_chunks[:chk])
airquali_text_formatted = format_llama3_prompt(airquali_text)
airquali_report = get_completion(airquali_text_formatted)

# %%
comfort_report = comfort_intro + thermal_report + daylight_report + acoustic_report + airquali_report

# %%
dl_means, dl_issues, dl_factor_issue_counts, dl_worst_factor = extract_delight_summary(del_path)
del_rag_chunks = create_rag_chunks(delight_reference)
keywords = ["view", "balcony", "space", "green", "window", "size"]
delight_chunks = [chunk for chunk in del_rag_chunks if any(kw in chunk.lower() for kw in keywords)]
delight_text = prompt_delight(dl_means, dl_issues, dl_factor_issue_counts, dl_worst_factor, delight_chunks[:chk])
delight_text_formatted = format_llama3_prompt(delight_text)
delight_report = get_completion(delight_text_formatted)

# %%
s_means, s_issues, s_factor_issue_counts, s_worst_factor = extract_social_summary(soc_path)
soc_rag_chunks = create_rag_chunks(social_reference)
keywords = ["social", "community", "interaction", "communal", "shared", "lobby", "green", "garden", "pocket", "buffer", "semi-private", "public space", "distribution"]
social_chunks = [chunk for chunk in soc_rag_chunks if any(kw in chunk.lower() for kw in keywords)]
social_text = prompt_social(s_means, s_issues, s_factor_issue_counts, s_worst_factor, social_chunks[:chk])
social_text_formatted = format_llama3_prompt(social_text)
social_report = get_completion(social_text_formatted)


collected_citation_keys = set()
collected_reference_extras = []  

def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    pattern = re.compile(r"(\*\*[^\*]+\*\*|https?://\S+|\S+|\s+)")

    for part in pattern.findall(text):
        if not part.strip():
            p.add_run(part)  # spaces
        elif part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("http://") or part.startswith("https://"):
            run = p.add_run(part)
            run.font.color.rgb = RGBColor(0, 102, 204)
            run.underline = True
            try:
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn as _qn

                r_id = doc.part.relate_to(
                    part,
                    reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                    is_external=True,
                )
                
                r = run._r
                hlink = OxmlElement("w:hyperlink")
                hlink.set(_qn("r:id"), r_id)
                r.addprevious(hlink)
                hlink.append(r)
            except Exception as e:
                print(f"⚠ Could not embed link for {part}:", e)
        else:
            p.add_run(part)

    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)

def add_report_section(doc, title, report_text, dimension, image_dir):
    """
    Ignores any '## References' printed by the model. We build the final References
    from in-text citations + MASTER_BIB to ensure correctness and completeness.
    """
    global collected_citation_keys, collected_reference_extras

    doc.add_page_break()
    doc.add_heading(title, level=1)

   
    collected_citation_keys |= extract_citation_keys(report_text)

    lines = [ln.strip() for ln in report_text.strip().split("\n") if ln.strip()]
    in_refs = False
    for line in lines:
        low = line.lower()

        
        if "doi:" in low or "https://doi.org" in low:
            collected_reference_extras.append(line)

        if f'{dimension}_images' in low:
            img1 = image_dir / f"{dimension}_factors.png"
            img2 = image_dir / f"{dimension}_satisfaction.png"
            table = doc.add_table(rows=1, cols=2)
            row = table.rows[0].cells
            if img1.exists():
                row[0].paragraphs[0].add_run().add_picture(str(img1), width=Inches(3))
            else:
                row[0].text = "[Missing image]"
            if img2.exists():
                row[1].paragraphs[0].add_run().add_picture(str(img2), width=Inches(3))
            else:
                row[1].text = "[Missing image]"
            doc.add_paragraph()
            continue

        if low.startswith("## references"):
            in_refs = True
            continue

        if line.startswith("## "):
            in_refs = False
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            add_body_paragraph(doc, "• " + line[2:])
        elif not in_refs:
            add_body_paragraph(doc, line)

def write_final_references(doc):
    refs = format_references_from_keys(collected_citation_keys, extras=collected_reference_extras)
    if not refs:
        return
    doc.add_page_break()
    doc.add_heading("References", level=1)
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)


doc = Document()
doc.add_heading("Wellbeing Fostered by Design Report", 0)

add_report_section(doc, "Wellbeing", wellbeing_report, "wellbeing", wb_im_path)
add_report_section(doc, "Comfort", comfort_report, "comfort", com_im_path)
add_report_section(doc, "Delight", delight_report, "delight", del_im_path)
add_report_section(doc, "Social", social_report, "social", soc_im_path)

write_final_references(doc)

output_path = out_f / "Wellbeing_Report.docx"
doc.save(output_path)
os.startfile(output_path)
